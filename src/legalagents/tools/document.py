"""
legalagents.tools.document
──────────────────────────
Tools de traitement documentaire — réécriture Python des capacités de mike.

Trois tools concrets (pas abstraits) — prêts à l'emploi :

  ReadDocumentTool       — Lecture PDF / DOCX → texte
  GenerateDocxTool       — Génération DOCX depuis contenu structuré
  TrackedChangesTool     — Modifications DOCX comme suivi Word (Accept/Reject)
  TabularAnalysisTool    — Analyse N documents × M critères → matrice

TrackedChangesTool est le bijou technique : il injecte des balises
<w:ins>/<w:del> nativement reconnues par Word et LibreOffice, sans
passer par une conversion ou un service externe.
"""

from __future__ import annotations

import copy
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from .base import LegalTool

# ── Namespaces XML Word ────────────────────────────────────────────────────────

W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
W     = f"{{{W_NS}}}"
XML   = f"{{{XML_NS}}}"

NSMAP = {
    "w":   W_NS,
    "xml": XML_NS,
}


# ── Dataclasses ───────────────────────────────────────────────────────────────

@dataclass
class EditInput:
    """Une modification à appliquer sur un document DOCX."""
    find: str
    replace: str
    context_before: str = ""
    context_after: str = ""
    reason: Optional[str] = None


@dataclass
class AppliedChange:
    """Résultat de l'application d'une modification."""
    change_id: str
    find: str
    replace: str
    reason: Optional[str]
    applied: bool
    paragraph_index: int = -1
    error: Optional[str] = None


@dataclass
class DocxSection:
    """Section d'un document généré."""
    heading: str
    content: str = ""
    level: int = 1                    # 1 = H1, 2 = H2
    table: Optional[dict] = None      # {"headers": [...], "rows": [[...]]}


@dataclass
class TabularCell:
    """Cellule d'une analyse tabulaire."""
    summary: str
    flag: Optional[str] = None        # "⚠️", "✅", "❌" ou None
    reasoning: str = ""


# ── Helpers XML ───────────────────────────────────────────────────────────────

def _set_preserve(t_elem) -> None:
    """Ajoute xml:space='preserve' si le texte commence/finit par un espace."""
    if t_elem.text and (t_elem.text.startswith(" ") or t_elem.text.endswith(" ")):
        t_elem.set(f"{XML}space", "preserve")


def _para_flat_text(para) -> str:
    """Texte plat d'un paragraphe (ignore le contenu w:del)."""
    parts = []
    for elem in para.iter():
        # Ignorer les textes supprimés
        if elem.tag == f"{W}delText":
            continue
        if elem.tag == f"{W}t" and elem.text:
            parts.append(elem.text)
    return "".join(parts)


def _para_char_map(para) -> tuple[str, list[tuple]]:
    """
    Construit le mapping char_idx → (run_elem, offset_in_run).
    Seuls les runs visibles (hors w:del) sont inclus.
    """
    flat: list[str] = []
    char_map: list[tuple] = []       # (run_elem, offset_in_run)

    def _collect(elem, in_del: bool = False):
        tag_local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag_local == "del":
            return                   # tout le contenu est ignoré
        if tag_local == "r":
            t = elem.find(f"{W}t")
            if t is not None and t.text:
                for i, ch in enumerate(t.text):
                    flat.append(ch)
                    char_map.append((elem, i))
            return
        for child in elem:
            _collect(child, in_del)

    for child in para:
        _collect(child)

    return "".join(flat), char_map


def _get_rpr(run_elem) -> Optional[Any]:
    """Récupère l'élément w:rPr (mise en forme) d'un run."""
    return run_elem.find(f"{W}rPr")


def _make_text_run(text: str, rpr=None) -> Any:
    """Crée un <w:r> avec le texte et la mise en forme donnés."""
    from lxml import etree
    r = etree.Element(f"{W}r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = etree.SubElement(r, f"{W}t")
    t.text = text
    _set_preserve(t)
    return r


def _make_del_run(text: str, rpr=None) -> Any:
    """Crée un <w:r> avec <w:delText> pour le contenu supprimé."""
    from lxml import etree
    r = etree.Element(f"{W}r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = etree.SubElement(r, f"{W}delText")
    t.text = text
    _set_preserve(t)
    return r


def _apply_edit_to_paragraph(
    para,
    edit: EditInput,
    del_id: int,
    ins_id: int,
    author: str,
    date_str: str,
) -> bool:
    """
    Applique une modification trackée à un paragraphe XML.

    Stratégie : on reconstruit le paragraphe en trois parties —
    texte avant, w:del (ancien), w:ins (nouveau), texte après.
    La mise en forme du run portant le texte cible est conservée.

    Returns True si la modification a été appliquée.
    """
    from lxml import etree

    flat, char_map = _para_char_map(para)
    if not flat:
        return False

    # Trouver la position cible avec le contexte comme ancrage
    ctx_before = edit.context_before or ""
    ctx_after  = edit.context_after  or ""
    find_text  = edit.find

    search = ctx_before + find_text + ctx_after
    pos = flat.find(search) if search.strip() else -1

    if pos != -1:
        target_start = pos + len(ctx_before)
    else:
        pos = flat.find(find_text)
        if pos == -1:
            return False
        target_start = pos

    target_end = target_start + len(find_text)

    if target_end > len(char_map):
        return False

    # Mise en forme depuis le premier run impliqué
    first_run = char_map[target_start][0]
    rpr = _get_rpr(first_run)

    # Propriétés du paragraphe (w:pPr) — à conserver
    ppr = para.find(f"{W}pPr")

    # Textes des trois zones
    before_text = flat[:target_start]
    target_text = flat[target_start:target_end]
    after_text  = flat[target_end:]

    # Vider le paragraphe (garder pPr)
    for child in list(para):
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local != "pPr":
            para.remove(child)

    # Reconstruire
    if before_text:
        para.append(_make_text_run(before_text, rpr))

    # w:del
    del_elem = etree.SubElement(para, f"{W}del", {
        f"{W}id":     str(del_id),
        f"{W}author": author,
        f"{W}date":   date_str,
    })
    del_elem.append(_make_del_run(target_text, rpr))

    # w:ins
    ins_elem = etree.SubElement(para, f"{W}ins", {
        f"{W}id":     str(ins_id),
        f"{W}author": author,
        f"{W}date":   date_str,
    })
    ins_run = etree.SubElement(ins_elem, f"{W}r")
    if rpr is not None:
        ins_run.append(copy.deepcopy(rpr))
    ins_t = etree.SubElement(ins_run, f"{W}t")
    ins_t.text = edit.replace
    _set_preserve(ins_t)

    if after_text:
        para.append(_make_text_run(after_text, rpr))

    return True


# ══════════════════════════════════════════════════════════════════════════════
# ReadDocumentTool
# ══════════════════════════════════════════════════════════════════════════════

class ReadDocumentTool(LegalTool):
    """
    Lit le contenu textuel d'un document PDF ou DOCX.

    Supporte :
      - PDF     : extraction avec pdfplumber (tableaux inclus)
      - DOCX    : extraction python-docx (paragraphes + tableaux)
      - TXT/MD  : lecture directe
    """

    name = "read_document"
    description = (
        "Lit le contenu textuel d'un document (PDF, DOCX, TXT). "
        "Retourne le texte extrait avec les tableaux formatés. "
        "Appeler avant toute analyse documentaire."
    )
    inputs = {
        "path": {
            "type": "string",
            "description": "Chemin absolu vers le fichier à lire",
        },
        "max_chars": {
            "type": "integer",
            "description": "Nombre max de caractères retournés (défaut 50 000)",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(self, path: str, max_chars: int = 50_000) -> str:
        p = Path(path)
        if not p.exists():
            return f"❌ Fichier introuvable : {path}"

        suffix = p.suffix.lower()
        try:
            if suffix == ".pdf":
                return self._read_pdf(p, max_chars)
            elif suffix in (".docx",):
                return self._read_docx(p, max_chars)
            elif suffix in (".txt", ".md"):
                return p.read_text(encoding="utf-8", errors="replace")[:max_chars]
            else:
                return f"❌ Format non supporté : {suffix}. Formats acceptés : PDF, DOCX, TXT, MD."
        except Exception as e:
            return f"❌ Erreur lecture document : {e}"

    def _read_pdf(self, path: Path, max_chars: int) -> str:
        try:
            import pdfplumber
        except ImportError:
            return self._read_pdf_fallback(path, max_chars)

        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")

                # Tableaux
                tables = page.extract_tables()
                for table in tables:
                    if not table:
                        continue
                    rows = [" | ".join(str(c or "") for c in row) for row in table]
                    text_parts.append("\n[TABLEAU]\n" + "\n".join(rows) + "\n[/TABLEAU]")

        result = "\n\n".join(text_parts)
        return result[:max_chars]

    def _read_pdf_fallback(self, path: Path, max_chars: int) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n\n".join(pages)[:max_chars]
        except ImportError:
            return "❌ pdfplumber ou pypdf requis. Installer : pip install pdfplumber"

    def _read_docx(self, path: Path, max_chars: int) -> str:
        from docx import Document
        doc = Document(str(path))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            parts.append("\n[TABLEAU]\n" + "\n".join(rows) + "\n[/TABLEAU]")

        return "\n\n".join(parts)[:max_chars]


# ══════════════════════════════════════════════════════════════════════════════
# GenerateDocxTool
# ══════════════════════════════════════════════════════════════════════════════

class GenerateDocxTool(LegalTool):
    """
    Génère un document DOCX structuré depuis du contenu JSON.

    Supporte : titre, sections avec niveaux de titres, tableaux, orientation.
    Le document est sauvegardé au chemin spécifié et son chemin est retourné.
    """

    name = "generate_docx"
    description = (
        "Génère un document Word (.docx) depuis un contenu structuré. "
        "Utiliser pour créer des synthèses, checklists, rapports d'analyse. "
        "Retourne le chemin du fichier généré."
    )
    inputs = {
        "title": {
            "type": "string",
            "description": "Titre du document",
        },
        "sections": {
            "type": "array",
            "description": (
                "Liste de sections. Chaque section est un objet avec : "
                "heading (str), content (str), level (int 1-3), "
                "table (dict avec headers et rows, optionnel)."
            ),
        },
        "output_path": {
            "type": "string",
            "description": "Chemin de sauvegarde du fichier .docx",
        },
        "landscape": {
            "type": "boolean",
            "description": "Orientation paysage (défaut : portrait)",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(
        self,
        title: str,
        sections: list[dict],
        output_path: str,
        landscape: bool = False,
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return "❌ python-docx requis : pip install python-docx"

        doc = Document()

        if landscape:
            section = doc.sections[0]
            section.orientation = 1  # WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        # Titre
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for sec in sections:
            heading  = sec.get("heading", "")
            content  = sec.get("content", "")
            level    = sec.get("level", 1)
            table    = sec.get("table")

            if heading:
                doc.add_heading(heading, level=level)

            if content:
                doc.add_paragraph(content)

            if table and isinstance(table, dict):
                headers = table.get("headers", [])
                rows    = table.get("rows", [])
                if headers:
                    t = doc.add_table(rows=1, cols=len(headers))
                    t.style = "Table Grid"
                    hdr_cells = t.rows[0].cells
                    for i, h in enumerate(headers):
                        hdr_cells[i].text = str(h)
                    for row_data in rows:
                        row_cells = t.add_row().cells
                        for i, val in enumerate(row_data[:len(headers)]):
                            row_cells[i].text = str(val)

        doc.save(output_path)
        return f"✅ Document généré : {output_path}"


# ══════════════════════════════════════════════════════════════════════════════
# TrackedChangesTool  ← LE BIJOU
# ══════════════════════════════════════════════════════════════════════════════

class TrackedChangesTool(LegalTool):
    """
    Propose des modifications sur un document .docx sous forme de suivi
    des modifications Word natif (Accept/Reject).

    Chaque modification produit une paire <w:del>/<w:ins> dans le XML
    du document, reconnue nativement par Word et LibreOffice.
    L'utilisateur peut accepter ou rejeter chaque changement individuellement.

    Inspiré de mike (mikeoss.com) — réécrit complètement en Python.
    Supérieur à la version TypeScript originale car :
      - Pas de dépendance externe (JSZip, fast-xml-parser)
      - Gestion native des namespaces XML Word
      - Contexte d'ancrage bidirectionnel (before + after)
      - Rapport détaillé par modification
    """

    name = "edit_document_tracked"
    description = (
        "Propose des modifications sur un document .docx comme suivi des modifications Word. "
        "Chaque changement apparaît en Accept/Reject dans Word ou LibreOffice. "
        "Préférer cet outil à generate_docx quand le document existe déjà : "
        "on ne régénère pas tout, on modifie chirurgicalement."
    )
    inputs = {
        "input_path": {
            "type": "string",
            "description": "Chemin du fichier .docx à modifier",
        },
        "edits": {
            "type": "array",
            "description": (
                "Liste de modifications. Chaque modification est un objet avec : "
                "find (str, texte à remplacer), "
                "replace (str, nouveau texte), "
                "context_before (str, texte précédant pour l'ancrage, recommandé), "
                "context_after (str, texte suivant pour l'ancrage, recommandé), "
                "reason (str, motif de la modification, optionnel)."
            ),
        },
        "output_path": {
            "type": "string",
            "description": "Chemin de sauvegarde. Si omis, crée un fichier _tracked.docx",
            "nullable": True,
        },
        "author": {
            "type": "string",
            "description": "Nom de l'auteur des modifications (défaut : 'LegalAgent')",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(
        self,
        input_path: str,
        edits: list[dict],
        output_path: Optional[str] = None,
        author: str = "LegalAgent",
    ) -> str:
        try:
            from lxml import etree
        except ImportError:
            return "❌ lxml requis : pip install lxml"

        src = Path(input_path)
        if not src.exists():
            return f"❌ Fichier introuvable : {input_path}"
        if src.suffix.lower() != ".docx":
            return f"❌ Format non supporté : {src.suffix}. Seul .docx est accepté."

        # Destination
        if not output_path:
            output_path = str(src.with_stem(src.stem + "_tracked"))

        # Timestamp ISO 8601 pour les attributs Word
        date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

        # Parser les edits
        parsed_edits: list[EditInput] = []
        for e in edits:
            if not e.get("find") or "replace" not in e:
                continue
            parsed_edits.append(EditInput(
                find           = e["find"],
                replace        = e["replace"],
                context_before = e.get("context_before", ""),
                context_after  = e.get("context_after", ""),
                reason         = e.get("reason"),
            ))

        if not parsed_edits:
            return "❌ Aucune modification valide fournie."

        # Lire et modifier le DOCX (ZIP)
        applied_changes: list[AppliedChange] = []
        change_counter = 1

        with zipfile.ZipFile(src, "r") as zin:
            zip_contents: dict[str, bytes] = {}
            for name in zin.namelist():
                zip_contents[name] = zin.read(name)

        # Trouver le nom du fichier document principal
        doc_xml_name = "word/document.xml"
        if doc_xml_name not in zip_contents:
            # Certains fichiers Windows utilisent des backslashes
            candidates = [n for n in zip_contents if n.lower().endswith("document.xml")]
            if not candidates:
                return "❌ Structure DOCX invalide : word/document.xml introuvable."
            doc_xml_name = candidates[0]

        # Parser le XML
        doc_xml = zip_contents[doc_xml_name]
        root = etree.fromstring(doc_xml)

        # Appliquer chaque modification sur les paragraphes
        paragraphs = list(root.iter(f"{W}p"))

        for edit in parsed_edits:
            applied = False
            para_idx = -1

            for idx, para in enumerate(paragraphs):
                del_id = change_counter
                ins_id = change_counter + 1
                change_counter += 2

                if _apply_edit_to_paragraph(para, edit, del_id, ins_id, author, date_str):
                    applied = True
                    para_idx = idx
                    break

            applied_changes.append(AppliedChange(
                change_id       = str(uuid.uuid4())[:8],
                find            = edit.find,
                replace         = edit.replace,
                reason          = edit.reason,
                applied         = applied,
                paragraph_index = para_idx,
                error           = None if applied else "Texte non trouvé dans le document",
            ))

        # Réécrire le XML modifié
        modified_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        zip_contents[doc_xml_name] = modified_xml

        # Écrire le nouveau fichier DOCX
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in zip_contents.items():
                zout.writestr(name, data)

        # Rapport
        n_ok  = sum(1 for c in applied_changes if c.applied)
        n_err = len(applied_changes) - n_ok

        lines = [f"✅ Document modifié : **{output_path}**"]
        lines.append(f"{n_ok}/{len(applied_changes)} modification(s) appliquée(s)")
        if n_err:
            lines.append(f"⚠️ {n_err} modification(s) non trouvée(s) :")

        for c in applied_changes:
            status = "✅" if c.applied else "❌"
            reason = f" _{c.reason}_" if c.reason else ""
            lines.append(f"  {status} `{c.find[:60]}` → `{c.replace[:60]}`{reason}")
            if not c.applied and c.error:
                lines.append(f"     → {c.error}")

        lines.append("\n💡 Ouvrir le fichier dans Word ou LibreOffice pour Accept/Reject les modifications.")
        return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TabularAnalysisTool
# ══════════════════════════════════════════════════════════════════════════════

class TabularAnalysisTool(LegalTool):
    """
    Analyse N documents selon M critères → matrice de synthèse.

    Inspiré de mike — adapté pour la due diligence juridique française.

    Chaque cellule (document × critère) est extraite par l'agent via
    l'outil read_document + une question ciblée.

    Usage typique : comparer des baux, des contrats-cadres, des CGV
    sur des critères standardisés (durée, résiliation, garanties…).
    """

    name = "tabular_analysis"
    description = (
        "Analyse plusieurs documents selon des critères définis, produit une matrice. "
        "Idéal pour la due diligence : comparer N contrats sur M points clés. "
        "Retourne un tableau Markdown + une synthèse des points d'attention."
    )
    inputs = {
        "documents": {
            "type": "array",
            "description": (
                "Liste de documents à analyser. Chaque document est un objet avec : "
                "path (str, chemin du fichier), label (str, nom court)."
            ),
        },
        "columns": {
            "type": "array",
            "description": (
                "Critères d'analyse. Chaque colonne est un objet avec : "
                "name (str, nom de la colonne), "
                "question (str, question précise à poser au document), "
                "flag_if (str, condition de signalement ⚠️, optionnel)."
            ),
        },
        "output_path": {
            "type": "string",
            "description": "Chemin DOCX de sortie (optionnel)",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(
        self,
        documents: list[dict],
        columns: list[dict],
        output_path: Optional[str] = None,
    ) -> str:
        if not documents or not columns:
            return "❌ Fournir au moins un document et un critère."

        reader = ReadDocumentTool()

        # Lire tous les documents
        doc_contents: dict[str, str] = {}
        for doc in documents:
            path  = doc.get("path", "")
            label = doc.get("label", Path(path).stem)
            content = reader.forward(path)
            doc_contents[label] = content

        # Construire la matrice
        matrix: dict[str, dict[str, TabularCell]] = {}
        flags: list[str] = []

        for doc in documents:
            label   = doc.get("label", Path(doc.get("path", "?")).stem)
            matrix[label] = {}
            text = doc_contents.get(label, "")

            for col in columns:
                col_name  = col.get("name", "?")
                question  = col.get("question", f"Quel est {col_name} dans ce document ?")
                flag_if   = col.get("flag_if", "")

                # Extraction simplifiée : chercher la réponse dans le texte
                # En production, ceci serait un appel LLM via le modèle de l'agent
                cell = self._extract_cell(text, question, col_name)

                if flag_if and flag_if.lower() in (cell.summary or "").lower():
                    cell.flag = "⚠️"
                    flags.append(f"{label} / {col_name} : {cell.summary[:100]}")

                matrix[label][col_name] = cell

        # Formater en Markdown
        col_names = [c.get("name", "?") for c in columns]
        header = "| Document | " + " | ".join(col_names) + " |"
        sep    = "|---|" + "---|" * len(col_names)
        rows   = [header, sep]

        for doc_label, cols in matrix.items():
            cells = []
            for cn in col_names:
                cell = cols.get(cn)
                if cell:
                    flag = cell.flag or ""
                    cells.append(f"{flag} {cell.summary[:80]}")
                else:
                    cells.append("—")
            rows.append(f"| {doc_label} | " + " | ".join(cells) + " |")

        result = "\n".join(rows)

        if flags:
            result += "\n\n**⚠️ Points d'attention :**\n"
            result += "\n".join(f"- {f}" for f in flags)

        if output_path:
            gen = GenerateDocxTool()
            sections = [{"heading": "Matrice d'analyse", "content": "", "level": 1,
                         "table": {
                             "headers": ["Document"] + col_names,
                             "rows": [
                                 [label] + [
                                     f"{matrix[label][cn].flag or ''} {matrix[label][cn].summary[:80]}"
                                     for cn in col_names
                                 ]
                                 for label in matrix
                             ],
                         }}]
            gen.forward(
                title="Analyse Comparative",
                sections=sections,
                output_path=output_path,
            )
            result += f"\n\n📄 Rapport DOCX : {output_path}"

        return result

    def _extract_cell(self, doc_text: str, question: str, col_name: str) -> TabularCell:
        """
        Extraction basique par recherche de mots-clés.
        En production, remplacer par un appel LLM.
        """
        # Recherche des N premiers mots du nom de colonne dans le texte
        keywords = col_name.lower().split()
        for line in doc_text.split("\n"):
            if any(kw in line.lower() for kw in keywords):
                return TabularCell(summary=line.strip()[:200])
        return TabularCell(summary="Non trouvé dans le document")
