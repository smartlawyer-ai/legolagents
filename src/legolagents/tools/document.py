"""
legolagents.tools.document
──────────────────────────
Document processing tools — clean-room Python rewrite of mike's capabilities.

Three concrete (non-abstract) tools — ready to use:

  ReadDocumentTool       — PDF / DOCX reading → text
  GenerateDocxTool       — DOCX generation from structured content
  TrackedChangesTool     — DOCX edits as native Word tracked changes (Accept/Reject)
  TabularAnalysisTool    — N documents × M criteria analysis → matrix

TrackedChangesTool is the technical centerpiece: it injects <w:ins>/<w:del>
tags natively recognized by Word and LibreOffice, without any conversion
or external service.
"""

from __future__ import annotations

import copy
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from .base import LegalTool

# ── Word XML namespaces ────────────────────────────────────────────────────────

W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
W     = f"{{{W_NS}}}"
XML   = f"{{{XML_NS}}}"

NSMAP = {
    "w":   W_NS,
    "xml": XML_NS,
}


# ── Dataclasses ───────────────────────────────────────────────────────────────

@dataclass
class EditInput:
    """A single edit to apply to a DOCX document."""
    find: str
    replace: str
    context_before: str = ""
    context_after: str = ""
    reason: Optional[str] = None


@dataclass
class AppliedChange:
    """Result of applying one edit."""
    change_id: str
    find: str
    replace: str
    reason: Optional[str]
    applied: bool
    paragraph_index: int = -1
    error: Optional[str] = None


@dataclass
class DocxSection:
    """Section of a generated document."""
    heading: str
    content: str = ""
    level: int = 1                    # 1 = H1, 2 = H2
    table: Optional[dict] = None      # {"headers": [...], "rows": [[...]]}


@dataclass
class TabularCell:
    """Cell of a tabular analysis."""
    summary: str
    flag: Optional[str] = None        # "⚠️", "✅", "❌" or None
    reasoning: str = ""


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_preserve(t_elem) -> None:
    """Add xml:space='preserve' if the text starts/ends with a space."""
    if t_elem.text and (t_elem.text.startswith(" ") or t_elem.text.endswith(" ")):
        t_elem.set(f"{XML}space", "preserve")


def _para_flat_text(para) -> str:
    """Flat text of a paragraph (ignores w:del content)."""
    parts = []
    for elem in para.iter():
        # Skip deleted text
        if elem.tag == f"{W}delText":
            continue
        if elem.tag == f"{W}t" and elem.text:
            parts.append(elem.text)
    return "".join(parts)


def _para_char_map(para) -> tuple[str, list[tuple]]:
    """
    Build the char_idx → (run_elem, offset_in_run) mapping.
    Only visible runs (outside w:del) are included.
    """
    flat: list[str] = []
    char_map: list[tuple] = []       # (run_elem, offset_in_run)

    def _collect(elem, in_del: bool = False):
        tag_local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag_local == "del":
            return                   # all content is ignored
        if tag_local == "r":
            t = elem.find(f"{W}t")
            if t is not None and t.text:
                for i, ch in enumerate(t.text):
                    flat.append(ch)
                    char_map.append((elem, i))
            return
        for child in elem:
            _collect(child, in_del)

    for child in para:
        _collect(child)

    return "".join(flat), char_map


def _get_rpr(run_elem) -> Optional[Any]:
    """Get the w:rPr (formatting) element of a run."""
    return run_elem.find(f"{W}rPr")


def _make_text_run(text: str, rpr=None) -> Any:
    """Create a <w:r> with the given text and formatting."""
    from lxml import etree
    r = etree.Element(f"{W}r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = etree.SubElement(r, f"{W}t")
    t.text = text
    _set_preserve(t)
    return r


def _make_del_run(text: str, rpr=None) -> Any:
    """Create a <w:r> with <w:delText> for deleted content."""
    from lxml import etree
    r = etree.Element(f"{W}r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = etree.SubElement(r, f"{W}delText")
    t.text = text
    _set_preserve(t)
    return r


def _apply_edit_to_paragraph(
    para,
    edit: EditInput,
    del_id: int,
    ins_id: int,
    author: str,
    date_str: str,
) -> bool:
    """
    Apply a tracked edit to an XML paragraph.

    Strategy: rebuild the paragraph in three parts —
    text before, w:del (old), w:ins (new), text after.
    The formatting of the run carrying the target text is preserved.

    Returns True if the edit was applied.
    """
    from lxml import etree

    flat, char_map = _para_char_map(para)
    if not flat:
        return False

    # Find the target position using the context as an anchor
    ctx_before = edit.context_before or ""
    ctx_after  = edit.context_after  or ""
    find_text  = edit.find

    search = ctx_before + find_text + ctx_after
    pos = flat.find(search) if search.strip() else -1

    if pos != -1:
        target_start = pos + len(ctx_before)
    else:
        pos = flat.find(find_text)
        if pos == -1:
            return False
        target_start = pos

    target_end = target_start + len(find_text)

    if target_end > len(char_map):
        return False

    # Formatting from the first run involved
    first_run = char_map[target_start][0]
    rpr = _get_rpr(first_run)

    # Paragraph properties (w:pPr) — to preserve
    ppr = para.find(f"{W}pPr")

    # Text of the three zones
    before_text = flat[:target_start]
    target_text = flat[target_start:target_end]
    after_text  = flat[target_end:]

    # Clear the paragraph (keep pPr)
    for child in list(para):
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local != "pPr":
            para.remove(child)

    # Rebuild
    if before_text:
        para.append(_make_text_run(before_text, rpr))

    # w:del
    del_elem = etree.SubElement(para, f"{W}del", {
        f"{W}id":     str(del_id),
        f"{W}author": author,
        f"{W}date":   date_str,
    })
    del_elem.append(_make_del_run(target_text, rpr))

    # w:ins
    ins_elem = etree.SubElement(para, f"{W}ins", {
        f"{W}id":     str(ins_id),
        f"{W}author": author,
        f"{W}date":   date_str,
    })
    ins_run = etree.SubElement(ins_elem, f"{W}r")
    if rpr is not None:
        ins_run.append(copy.deepcopy(rpr))
    ins_t = etree.SubElement(ins_run, f"{W}t")
    ins_t.text = edit.replace
    _set_preserve(ins_t)

    if after_text:
        para.append(_make_text_run(after_text, rpr))

    return True


# ══════════════════════════════════════════════════════════════════════════════
# ReadDocumentTool
# ══════════════════════════════════════════════════════════════════════════════

class ReadDocumentTool(LegalTool):
    """
    Reads the text content of a PDF or DOCX document.

    Supports:
      - PDF     : extraction with pdfplumber (tables included)
      - DOCX    : extraction with python-docx (paragraphs + tables)
      - TXT/MD  : direct reading
    """

    name = "read_document"
    description = (
        "Reads the text content of a document (PDF, DOCX, TXT). "
        "Returns the extracted text with tables formatted. "
        "Call before any document analysis."
    )
    inputs = {
        "path": {
            "type": "string",
            "description": "Absolute path to the file to read",
        },
        "max_chars": {
            "type": "integer",
            "description": "Max number of characters returned (default 50,000)",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(self, path: str, max_chars: int = 50_000) -> str:
        p = Path(path)
        if not p.exists():
            return f"❌ File not found: {path}"

        suffix = p.suffix.lower()
        try:
            if suffix == ".pdf":
                return self._read_pdf(p, max_chars)
            elif suffix in (".docx",):
                return self._read_docx(p, max_chars)
            elif suffix in (".txt", ".md"):
                return p.read_text(encoding="utf-8", errors="replace")[:max_chars]
            else:
                return f"❌ Unsupported format: {suffix}. Accepted formats: PDF, DOCX, TXT, MD."
        except Exception as e:
            return f"❌ Error reading document: {e}"

    def _read_pdf(self, path: Path, max_chars: int) -> str:
        try:
            import pdfplumber
        except ImportError:
            return self._read_pdf_fallback(path, max_chars)

        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")

                # Tables
                tables = page.extract_tables()
                for table in tables:
                    if not table:
                        continue
                    rows = [" | ".join(str(c or "") for c in row) for row in table]
                    text_parts.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")

        result = "\n\n".join(text_parts)
        return result[:max_chars]

    def _read_pdf_fallback(self, path: Path, max_chars: int) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n\n".join(pages)[:max_chars]
        except ImportError:
            return "❌ pdfplumber or pypdf required. Install: pip install pdfplumber"

    def _read_docx(self, path: Path, max_chars: int) -> str:
        from docx import Document
        doc = Document(str(path))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            parts.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")

        return "\n\n".join(parts)[:max_chars]


# ══════════════════════════════════════════════════════════════════════════════
# GenerateDocxTool
# ══════════════════════════════════════════════════════════════════════════════

class GenerateDocxTool(LegalTool):
    """
    Generates a structured DOCX document from JSON content.

    Supports: title, sections with heading levels, tables, orientation.
    The document is saved to the given path, which is then returned.
    """

    name = "generate_docx"
    description = (
        "Generates a Word (.docx) document from structured content. "
        "Use to create summaries, checklists, analysis reports. "
        "Returns the path of the generated file."
    )
    inputs = {
        "title": {
            "type": "string",
            "description": "Document title",
        },
        "sections": {
            "type": "array",
            "description": (
                "List of sections. Each section is an object with: "
                "heading (str), content (str), level (int 1-3), "
                "table (dict with headers and rows, optional)."
            ),
        },
        "output_path": {
            "type": "string",
            "description": "Path to save the .docx file",
        },
        "landscape": {
            "type": "boolean",
            "description": "Landscape orientation (default: portrait)",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(
        self,
        title: str,
        sections: list[dict],
        output_path: str,
        landscape: bool = False,
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return "❌ python-docx required: pip install python-docx"

        doc = Document()

        if landscape:
            section = doc.sections[0]
            section.orientation = 1  # WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for sec in sections:
            heading  = sec.get("heading", "")
            content  = sec.get("content", "")
            level    = sec.get("level", 1)
            table    = sec.get("table")

            if heading:
                doc.add_heading(heading, level=level)

            if content:
                doc.add_paragraph(content)

            if table and isinstance(table, dict):
                headers = table.get("headers", [])
                rows    = table.get("rows", [])
                if headers:
                    t = doc.add_table(rows=1, cols=len(headers))
                    t.style = "Table Grid"
                    hdr_cells = t.rows[0].cells
                    for i, h in enumerate(headers):
                        hdr_cells[i].text = str(h)
                    for row_data in rows:
                        row_cells = t.add_row().cells
                        for i, val in enumerate(row_data[:len(headers)]):
                            row_cells[i].text = str(val)

        doc.save(output_path)
        return f"✅ Document generated: {output_path}"


# ══════════════════════════════════════════════════════════════════════════════
# TrackedChangesTool  ← THE CENTERPIECE
# ══════════════════════════════════════════════════════════════════════════════

class TrackedChangesTool(LegalTool):
    """
    Proposes edits to a .docx document as native Word tracked changes
    (Accept/Reject).

    Each edit produces a <w:del>/<w:ins> pair in the document's XML,
    natively recognized by Word and LibreOffice. The user can accept or
    reject each change individually.

    Inspired by mike (mikeoss.com) — fully rewritten in Python.
    Superior to the original TypeScript version because:
      - No external dependency (JSZip, fast-xml-parser)
      - Native handling of Word XML namespaces
      - Bidirectional anchoring context (before + after)
      - Detailed report per edit
    """

    name = "edit_document_tracked"
    description = (
        "Proposes edits to a .docx document as Word tracked changes. "
        "Each change appears as Accept/Reject in Word or LibreOffice. "
        "Prefer this tool over generate_docx when the document already exists: "
        "don't regenerate everything, edit surgically."
    )
    inputs = {
        "input_path": {
            "type": "string",
            "description": "Path of the .docx file to edit",
        },
        "edits": {
            "type": "array",
            "description": (
                "List of edits. Each edit is an object with: "
                "find (str, text to replace), "
                "replace (str, new text), "
                "context_before (str, preceding text for anchoring, recommended), "
                "context_after (str, following text for anchoring, recommended), "
                "reason (str, reason for the edit, optional)."
            ),
        },
        "output_path": {
            "type": "string",
            "description": "Save path. If omitted, creates a _tracked.docx file",
            "nullable": True,
        },
        "author": {
            "type": "string",
            "description": "Name of the change author (default: 'LegalAgent')",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(
        self,
        input_path: str,
        edits: list[dict],
        output_path: Optional[str] = None,
        author: str = "LegalAgent",
    ) -> str:
        try:
            from lxml import etree
        except ImportError:
            return "❌ lxml required: pip install lxml"

        src = Path(input_path)
        if not src.exists():
            return f"❌ File not found: {input_path}"
        if src.suffix.lower() != ".docx":
            return f"❌ Unsupported format: {src.suffix}. Only .docx is accepted."

        # Destination
        if not output_path:
            output_path = str(src.with_stem(src.stem + "_tracked"))

        # ISO 8601 timestamp for Word attributes
        date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

        # Parse the edits
        parsed_edits: list[EditInput] = []
        for e in edits:
            if not e.get("find") or "replace" not in e:
                continue
            parsed_edits.append(EditInput(
                find           = e["find"],
                replace        = e["replace"],
                context_before = e.get("context_before", ""),
                context_after  = e.get("context_after", ""),
                reason         = e.get("reason"),
            ))

        if not parsed_edits:
            return "❌ No valid edit provided."

        # Read and edit the DOCX (ZIP)
        applied_changes: list[AppliedChange] = []
        change_counter = 1

        with zipfile.ZipFile(src, "r") as zin:
            zip_contents: dict[str, bytes] = {}
            for name in zin.namelist():
                zip_contents[name] = zin.read(name)

        # Find the main document file name
        doc_xml_name = "word/document.xml"
        if doc_xml_name not in zip_contents:
            # Some Windows-produced files use backslashes
            candidates = [n for n in zip_contents if n.lower().endswith("document.xml")]
            if not candidates:
                return "❌ Invalid DOCX structure: word/document.xml not found."
            doc_xml_name = candidates[0]

        # Parse the XML
        doc_xml = zip_contents[doc_xml_name]
        root = etree.fromstring(doc_xml)

        # Apply each edit to the paragraphs
        paragraphs = list(root.iter(f"{W}p"))

        for edit in parsed_edits:
            applied = False
            para_idx = -1

            for idx, para in enumerate(paragraphs):
                del_id = change_counter
                ins_id = change_counter + 1
                change_counter += 2

                if _apply_edit_to_paragraph(para, edit, del_id, ins_id, author, date_str):
                    applied = True
                    para_idx = idx
                    break

            applied_changes.append(AppliedChange(
                change_id       = str(uuid.uuid4())[:8],
                find            = edit.find,
                replace         = edit.replace,
                reason          = edit.reason,
                applied         = applied,
                paragraph_index = para_idx,
                error           = None if applied else "Text not found in the document",
            ))

        # Rewrite the modified XML
        modified_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        zip_contents[doc_xml_name] = modified_xml

        # Write the new DOCX file
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in zip_contents.items():
                zout.writestr(name, data)

        # Report
        n_ok  = sum(1 for c in applied_changes if c.applied)
        n_err = len(applied_changes) - n_ok

        lines = [f"✅ Document edited: **{output_path}**"]
        lines.append(f"{n_ok}/{len(applied_changes)} edit(s) applied")
        if n_err:
            lines.append(f"⚠️ {n_err} edit(s) not found:")

        for c in applied_changes:
            status = "✅" if c.applied else "❌"
            reason = f" _{c.reason}_" if c.reason else ""
            lines.append(f"  {status} `{c.find[:60]}` → `{c.replace[:60]}`{reason}")
            if not c.applied and c.error:
                lines.append(f"     → {c.error}")

        lines.append("\n💡 Open the file in Word or LibreOffice to Accept/Reject the changes.")
        return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# TabularAnalysisTool
# ══════════════════════════════════════════════════════════════════════════════

class TabularAnalysisTool(LegalTool):
    """
    Analyzes N documents against M criteria → summary matrix.

    Inspired by mike — adapted for legal due diligence, any jurisdiction.

    Each cell (document × criterion) is extracted by the agent via the
    read_document tool + a targeted question.

    Typical usage: compare leases, master agreements, terms and conditions
    against standardized criteria (term, termination, guarantees…).
    """

    name = "tabular_analysis"
    description = (
        "Analyzes multiple documents against defined criteria, produces a matrix. "
        "Ideal for due diligence: comparing N contracts on M key points. "
        "Returns a Markdown table + a summary of points of attention."
    )
    inputs = {
        "documents": {
            "type": "array",
            "description": (
                "List of documents to analyze. Each document is an object with: "
                "path (str, file path), label (str, short name)."
            ),
        },
        "columns": {
            "type": "array",
            "description": (
                "Analysis criteria. Each column is an object with: "
                "name (str, column name), "
                "question (str, precise question to ask of the document), "
                "flag_if (str, flagging condition ⚠️, optional)."
            ),
        },
        "output_path": {
            "type": "string",
            "description": "Output DOCX path (optional)",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(
        self,
        documents: list[dict],
        columns: list[dict],
        output_path: Optional[str] = None,
    ) -> str:
        if not documents or not columns:
            return "❌ Provide at least one document and one criterion."

        reader = ReadDocumentTool()

        # Read all documents
        doc_contents: dict[str, str] = {}
        for doc in documents:
            path  = doc.get("path", "")
            label = doc.get("label", Path(path).stem)
            content = reader.forward(path)
            doc_contents[label] = content

        # Build the matrix
        matrix: dict[str, dict[str, TabularCell]] = {}
        flags: list[str] = []

        for doc in documents:
            label   = doc.get("label", Path(doc.get("path", "?")).stem)
            matrix[label] = {}
            text = doc_contents.get(label, "")

            for col in columns:
                col_name  = col.get("name", "?")
                question  = col.get("question", f"What is {col_name} in this document?")
                flag_if   = col.get("flag_if", "")

                # Simplified extraction: look for the answer in the text
                # In production, this would be an LLM call via the agent's model
                cell = self._extract_cell(text, question, col_name)

                if flag_if and flag_if.lower() in (cell.summary or "").lower():
                    cell.flag = "⚠️"
                    flags.append(f"{label} / {col_name}: {cell.summary[:100]}")

                matrix[label][col_name] = cell

        # Format as Markdown
        col_names = [c.get("name", "?") for c in columns]
        header = "| Document | " + " | ".join(col_names) + " |"
        sep    = "|---|" + "---|" * len(col_names)
        rows   = [header, sep]

        for doc_label, cols in matrix.items():
            cells = []
            for cn in col_names:
                cell = cols.get(cn)
                if cell:
                    flag = cell.flag or ""
                    cells.append(f"{flag} {cell.summary[:80]}")
                else:
                    cells.append("—")
            rows.append(f"| {doc_label} | " + " | ".join(cells) + " |")

        result = "\n".join(rows)

        if flags:
            result += "\n\n**⚠️ Points of attention:**\n"
            result += "\n".join(f"- {f}" for f in flags)

        if output_path:
            gen = GenerateDocxTool()
            sections = [{"heading": "Analysis Matrix", "content": "", "level": 1,
                         "table": {
                             "headers": ["Document"] + col_names,
                             "rows": [
                                 [label] + [
                                     f"{matrix[label][cn].flag or ''} {matrix[label][cn].summary[:80]}"
                                     for cn in col_names
                                 ]
                                 for label in matrix
                             ],
                         }}]
            gen.forward(
                title="Comparative Analysis",
                sections=sections,
                output_path=output_path,
            )
            result += f"\n\n📄 DOCX report: {output_path}"

        return result

    def _extract_cell(self, doc_text: str, question: str, col_name: str) -> TabularCell:
        """
        Basic extraction via keyword search.
        In production, replace with an LLM call.
        """
        # Search for the column name's keywords in the text
        keywords = col_name.lower().split()
        for line in doc_text.split("\n"):
            if any(kw in line.lower() for kw in keywords):
                return TabularCell(summary=line.strip()[:200])
        return TabularCell(summary="Not found in the document")
